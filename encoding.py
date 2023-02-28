from docx import Document
from lxml import etree
import os
import random

# Set the paths for the input and output directories
input_dir = 'word'
output_dir = 'html'

# Create the output directory if it doesn't exist
if not os.path.exists(output_dir):
    os.makedirs(output_dir)

# Initialize the list of processed file names
processed_files = []

# Iterate through each file in the input directory
for filename in os.listdir(input_dir):
    # Check if the file is a Word document
    if filename.endswith('.docx') and not filename.startswith('~$'):
        # Add the file name to the list of processed file names
        processed_files.append(filename)

        # Open the Word document
        doc_path = os.path.join(input_dir, filename)
        doc = Document(doc_path)

        # Create an XML element for the HTML document
        html = etree.Element('html')
        body = etree.SubElement(html, 'body')

        # Initialize the variable to check for blank lines
        prev_text = ''

        # Iterate through each paragraph in the document
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():
                # Check if the paragraph starts with "This document" or "Last Modified"
                if paragraph.text.strip().startswith('This document') or paragraph.text.strip().startswith('Last Modified'):
                    # Skip the paragraph
                    continue

                # Create a new HTML paragraph element and add the text
                p = etree.Element('p')

                # Add a self-closing <a> tag at the start of the paragraph
                a_id = str(random.randint(100000000, 999999999))
                a = etree.Element('a', {'class': 'brl-location', 'id': a_id})
                a.tail = paragraph.text.strip()
                p.append(a)

                # Check if the next paragraph is blank and if the current paragraph is not the last one
                if i < len(doc.paragraphs) - 1 and not doc.paragraphs[i+1].text.strip():
                    p.set('class', 'brl-btmmargin')

                # Check if the paragraph is centered and add the appropriate class
                alignment = paragraph.alignment
                if alignment == 1:
                    if 'brl-btmmargin' not in p.attrib:
                        p.set('class', 'brl-btmmargin')
                    p.set('class', p.attrib.get('class', '') + ' brl-align-center')

                body.append(p)

                # Update the variable to check for blank lines
                prev_text = paragraph.text

        # Remove the "brl-btmmargin" class from the last paragraph, if present
        if len(body) > 0:
            body[-1].attrib.pop('class', None)

        # Serialize the HTML element to a string
        html_string = etree.tostring(html, encoding='unicode', pretty_print=True)

        # Create the output file path
        html_path = os.path.join(output_dir, os.path.splitext(filename)[0] + '.html')

        # Save the HTML string to a file
        with open(html_path, 'w') as file:
            file.write(html_string)

# Write the processed file names to a text file
with open(os.path.join(output_dir, 'mappings.txt'), 'w') as file:
    for filename in processed_files:
        file.write(filename + '\n')
